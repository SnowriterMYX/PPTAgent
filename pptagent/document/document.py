import asyncio
import re
import traceback
from dataclasses import asdict, dataclass
from datetime import datetime
from typing import Any, Optional
import os
import docx
import chardet

from jinja2 import Environment, StrictUndefined
from torch import cosine_similarity

from pptagent.agent import Agent, AsyncAgent
from pptagent.llms import LLM, AsyncLLM
from pptagent.utils import edit_distance, get_logger, package_join, pexists

from .element import Section, SubSection, Table, link_medias

logger = get_logger(__name__)

env = Environment(undefined=StrictUndefined)

MERGE_METADATA_PROMPT = env.from_string(
    open(package_join("prompts", "merge_metadata.txt"), encoding='utf-8').read()
)
HEADING_EXTRACT_PROMPT = env.from_string(
    open(package_join("prompts", "heading_extract.txt"), encoding='utf-8').read()
)
SECTION_SUMMARY_PROMPT = env.from_string(
    open(package_join("prompts", "section_summary.txt"), encoding='utf-8').read()
)

MARKDOWN_IMAGE_REGEX = re.compile(r"!\[.*\]\(.*\)")
MARKDOWN_TABLE_REGEX = re.compile(r"\|.*\|")


def split_markdown_by_headings(
    markdown_content: str,
    headings: list[str],
    adjusted_headings: list[str],
    min_chunk_size: int = 64,
) -> list[str]:
    """
    Split markdown content using headings as separators without regex.

    Args:
        markdown_content (str): The markdown content to split
        headings (list[str]): List of heading strings to split by

    Returns:
        list[str]: List of content sections
    """
    adjusted_headings = [
        max(headings, key=lambda x: edit_distance(x, ah)) for ah in adjusted_headings
    ]
    sections = []
    current_section = []

    for line in markdown_content.splitlines():
        if any(line.strip().startswith(h) for h in adjusted_headings):
            if len(current_section) != 0:
                sections.append("\n".join(current_section).strip())
            current_section = [line]
        else:
            current_section.append(line)

    if len(current_section) != 0:
        sections.append("\n".join(current_section).strip())

    # if an chunk is too small, merge it with the previous chunk
    for i in reversed(range(1, len(sections))):
        if len(sections[i]) < min_chunk_size:
            sections[i - 1] += sections[i]
            sections.pop(i)

    if len(sections) > 1 and len(sections[0]) < min_chunk_size:
        sections[0] += sections[1]
        sections.pop(1)

    return sections


def to_paragraphs(original_text: str, max_chunk_size: int = 256):
    paragraphs = []
    medias = []
    for i, para in enumerate(original_text.split("\n\n")):
        para = para.strip()
        if not para:
            continue
        paragraph = {"markdown_content": para, "index": i}
        if MARKDOWN_TABLE_REGEX.match(para):
            paragraph["type"] = "table"
            medias.append(paragraph)
        elif MARKDOWN_IMAGE_REGEX.match(para):
            paragraph["type"] = "image"
            medias.append(paragraph)
        else:
            paragraphs.append(paragraph)
    for media in medias:
        pre_chunk = ""
        after_chunk = ""
        for chunk in reversed(paragraphs):
            if chunk["index"] < media["index"]:
                pre_chunk += chunk["markdown_content"] + "\n\n"
                if len(pre_chunk) > max_chunk_size:
                    break
        for chunk in paragraphs:
            if chunk["index"] > media["index"]:
                after_chunk += chunk["markdown_content"] + "\n\n"
                if len(after_chunk) > max_chunk_size:
                    break
        media["near_chunks"] = (pre_chunk, after_chunk)
    return medias


@dataclass
class Document:
    image_dir: str
    sections: list[Section]
    metadata: dict[str, str]

    def __post_init__(self):
        self.metadata["presentation-date"] = datetime.now().strftime("%Y-%m-%d")

    def iter_medias(self):
        for section in self.sections:
            yield from section.iter_medias()

    def get_table(self, image_path: str):
        for media in self.iter_medias():
            if media.path == image_path and isinstance(media, Table):
                return media
        raise ValueError(f"table not found: {image_path}")

    @classmethod
    def from_dict(
        cls, data: dict[str, Any], image_dir: str, require_caption: bool = True
    ):
        assert (
            "sections" in data
        ), f"'sections' key is required in data dictionary but was not found. Input keys: {list(data.keys())}"
        assert (
            "metadata" in data
        ), f"'metadata' key is required in data dictionary but was not found. Input keys: {list(data.keys())}"
        assert pexists(image_dir), f"image directory is not found: {image_dir}"
        document = cls(
            image_dir=image_dir,
            sections=[Section.from_dict(section) for section in data["sections"]],
            metadata=data["metadata"],
        )
        for section in document.sections:
            section.validate_medias(image_dir, require_caption)
        return document

    @classmethod
    def _parse_chunk(
        cls,
        extractor: Agent,
        language_model: LLM,
        vision_model: LLM,
        table_model: LLM,
        metadata: Optional[dict[str, Any]],
        section: Optional[dict[str, Any]],
        image_dir: str,
        turn_id: int = None,
        retry: int = 0,
        medias: Optional[list[dict]] = None,
    ):
        if retry == 0:
            medias = to_paragraphs(section)
            turn_id, section = extractor(markdown_document=section)
            metadata = section.pop("metadata", {})
        try:
            section["subsections"] = link_medias(medias, section["subsections"])
            section = Section.from_dict(section)
            for media in section.iter_medias():
                media.parse(table_model, image_dir)
                if isinstance(media, Table):
                    media.get_caption(language_model)
                else:
                    media.get_caption(vision_model)
            section.validate_medias(image_dir, False)
        except Exception as e:
            if retry < 3:
                logger.info("Retry section with error: %s", str(e))
                new_section = extractor.retry(
                    str(e), traceback.format_exc(), turn_id, retry + 1
                )
                return cls._parse_chunk(
                    extractor,
                    language_model,
                    vision_model,
                    table_model,
                    metadata,
                    new_section,
                    image_dir,
                    turn_id,
                    retry + 1,
                    medias,
                )
            else:
                logger.error(
                    "Failed to extract section, tried %d times",
                    retry,
                    exc_info=e,
                )
                raise e
        return metadata, section

    @classmethod
    async def _parse_chunk_async(
        cls,
        extractor: AsyncAgent,
        language_model: AsyncLLM,
        vision_model: AsyncLLM,
        table_model: Optional[AsyncLLM],
        metadata: Optional[dict[str, Any]],
        section: Optional[dict[str, Any]],
        image_dir: str,
        turn_id: int = None,
        retry: int = 0,
        medias: Optional[list[dict]] = None,
    ):
        if retry == 0:
            medias = to_paragraphs(section)
            turn_id, section = await extractor(markdown_document=section)
            metadata = section.pop("metadata", {})
        try:
            section["subsections"] = link_medias(medias, section["subsections"])
            section = Section.from_dict(section)
            for media in section.iter_medias():
                await media.parse_async(table_model, image_dir)
                if isinstance(media, Table):
                    await media.get_caption_async(language_model)
                else:
                    await media.get_caption_async(vision_model)
            section.validate_medias(image_dir, False)
        except Exception as e:
            if retry < 3:
                logger.info("Retry section with error: %s", str(e))
                logger.debug("Section data that caused error: %s", section)
                new_section = await extractor.retry(
                    str(e), traceback.format_exc(), turn_id, retry + 1
                )
                logger.debug("Retry generated new section: %s", new_section)
                return await cls._parse_chunk_async(
                    extractor,
                    language_model,
                    vision_model,
                    table_model,
                    metadata,
                    new_section,
                    image_dir,
                    turn_id,
                    retry + 1,
                    medias,
                )
            else:
                logger.error(
                    "Failed to extract section, tried %d times",
                    retry,
                    exc_info=e,
                )
                raise e
        return metadata, section

    @classmethod
    def from_markdown(
        cls,
        markdown_content: str,
        language_model: LLM,
        vision_model: LLM,
        image_dir: str,
        table_model: Optional[LLM] = None,
        topic: Optional[str] = None,
    ):
        """
        Create a Document from markdown content.

        Args:
            markdown_content (str): The markdown content.
            language_model (LLM): The language model.
            vision_model (LLM): The vision model.
            image_dir (str): The directory containing images.
            topic (Optional[str]): The user-provided topic.

        Returns:
            Document: The created document.
        """
        doc_extractor = Agent(
            "doc_extractor",
            llm_mapping={"language": language_model, "vision": vision_model},
        )

        metadata_list = []
        sections = []

        headings = re.findall(r"^#+\s+.*", markdown_content, re.MULTILINE)
        adjusted_headings = language_model(
            HEADING_EXTRACT_PROMPT.render(headings=headings), return_json=True
        )

        for chunk in split_markdown_by_headings(
            markdown_content, headings, adjusted_headings
        ):
            metadata, section = cls._parse_chunk(
                doc_extractor,
                language_model,
                vision_model,
                table_model,
                None,
                chunk,
                image_dir,
            )
            section.summary = language_model(
                SECTION_SUMMARY_PROMPT.render(section_content=chunk),
            )
            metadata_list.append(metadata)
            sections.append(section)

        merged_metadata = language_model(
            MERGE_METADATA_PROMPT.render(metadata=metadata_list, topic=topic),
            return_json=True,
        )

        # 确保merged_metadata是字典类型
        if not isinstance(merged_metadata, dict):
            logger.warning(f"merged_metadata is not a dict: {type(merged_metadata)}, using default metadata")
            merged_metadata = {
                "title": "未知标题",
                "author": "未知作者",
                "subject": "演示文稿"
            }

        return Document(
            image_dir=image_dir, metadata=merged_metadata, sections=sections
        )

    @classmethod
    def from_text_file(
        cls,
        file_path: str,
        language_model: LLM,
        vision_model: LLM,
        image_dir: str,
        table_model: Optional[LLM] = None,
        file_type: str = "auto",
    ):
        """
        从各种文本文件格式创建Document对象

        Args:
            file_path (str): 文件路径
            language_model (LLM): 语言模型
            vision_model (LLM): 视觉模型
            image_dir (str): 图片目录
            table_model (Optional[LLM]): 表格模型
            file_type (str): 文件类型，支持 "auto", "txt", "md", "docx", "rtf"

        Returns:
            Document: 创建的文档对象
        """
        # 自动检测文件类型
        if file_type == "auto":
            file_type = cls._detect_file_type(file_path)

        # 根据文件类型解析内容
        if file_type == "txt":
            content = cls._parse_txt_file(file_path)
        elif file_type == "md" or file_type == "markdown":
            content = cls._parse_markdown_file(file_path)
        elif file_type == "docx":
            content = cls._parse_docx_file(file_path)
        elif file_type == "rtf":
            content = cls._parse_rtf_file(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

        # 使用from_markdown方法处理内容
        return cls.from_markdown(content, language_model, vision_model, image_dir, table_model)

    @classmethod
    def from_user_input(
        cls,
        user_text: str,
        language_model: LLM,
        vision_model: LLM,
        image_dir: str,
        table_model: Optional[LLM] = None,
        title: str = "用户输入文档",
    ):
        """
        从用户直接输入的文本创建Document对象

        Args:
            user_text (str): 用户输入的文本内容
            language_model (LLM): 语言模型
            vision_model (LLM): 视觉模型
            image_dir (str): 图片目录
            table_model (Optional[LLM]): 表格模型
            title (str): 文档标题

        Returns:
            Document: 创建的文档对象
        """
        # 将用户输入转换为markdown格式
        markdown_content = cls._format_user_input_to_markdown(user_text, title)

        # 使用from_markdown方法处理内容
        return cls.from_markdown(
            markdown_content, language_model, vision_model, image_dir, table_model, title
        )

    @staticmethod
    def _detect_file_type(file_path: str) -> str:
        """自动检测文件类型"""
        _, ext = os.path.splitext(file_path.lower())

        type_mapping = {
            '.txt': 'txt',
            '.md': 'md',
            '.markdown': 'md',
            '.docx': 'docx',
            '.rtf': 'rtf',
        }

        return type_mapping.get(ext, 'txt')

    @staticmethod
    def _parse_txt_file(file_path: str) -> str:
        """解析TXT文件"""
        try:
            # 尝试检测文件编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

            # 使用检测到的编码读取文件
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()

            return content
        except Exception as e:
            logger.warning(f"TXT文件解析失败，尝试使用UTF-8编码: {e}")
            # 备用方案：使用UTF-8编码
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    @staticmethod
    def _parse_markdown_file(file_path: str) -> str:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()

    @staticmethod
    def _parse_docx_file(file_path: str) -> str:
        """解析DOCX文件"""
        try:
            doc = docx.Document(file_path)
            content_parts = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 根据段落样式判断是否为标题
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        if level.isdigit():
                            content_parts.append(f"{'#' * int(level)} {text}")
                        else:
                            content_parts.append(f"## {text}")
                    else:
                        content_parts.append(text)
                    content_parts.append("")  # 添加空行

            return "\n".join(content_parts)
        except Exception as e:
            logger.error(f"DOCX文件解析失败: {e}")
            raise ValueError(f"无法解析DOCX文件: {e}")

    @staticmethod
    def _parse_rtf_file(file_path: str) -> str:
        """解析RTF文件（简单实现）"""
        try:
            # 简单的RTF解析，去除RTF控制字符
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # 移除RTF控制字符（简单处理）
            import re
            # 移除RTF头部
            content = re.sub(r'\\rtf\d+.*?\\deff\d+', '', content)
            # 移除控制字符
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            # 移除花括号
            content = re.sub(r'[{}]', '', content)
            # 清理多余空白
            content = re.sub(r'\s+', ' ', content).strip()

            return content
        except Exception as e:
            logger.error(f"RTF文件解析失败: {e}")
            raise ValueError(f"无法解析RTF文件: {e}")

    @staticmethod
    def _format_user_input_to_markdown(user_text: str, title: str) -> str:
        """将用户输入格式化为Markdown"""
        lines = user_text.strip().split('\n')
        formatted_lines = [f"# {title}", ""]

        current_section = []
        for line in lines:
            line = line.strip()
            if not line:
                if current_section:
                    formatted_lines.extend(current_section)
                    formatted_lines.append("")
                    current_section = []
                continue

            # 检测是否可能是标题（简单启发式）
            if (len(line) < 50 and
                not line.endswith('.') and
                not line.endswith(',') and
                not line.endswith(';')):
                # 可能是标题
                if current_section:
                    formatted_lines.extend(current_section)
                    formatted_lines.append("")
                    current_section = []
                formatted_lines.append(f"## {line}")
                formatted_lines.append("")
            else:
                current_section.append(line)

        # 添加剩余内容
        if current_section:
            formatted_lines.extend(current_section)

        return "\n".join(formatted_lines)

    @classmethod
    async def from_markdown_async(
        cls,
        markdown_content: str,
        language_model: AsyncLLM,
        vision_model: AsyncLLM,
        image_dir: str,
        table_model: Optional[AsyncLLM] = None,
        topic: Optional[str] = None,
    ):
        doc_extractor = AsyncAgent(
            "doc_extractor",
            llm_mapping={"language": language_model, "vision": vision_model},
        )

        headings = re.findall(r"^#+\s+.*", markdown_content, re.MULTILINE)
        adjusted_headings = await language_model(
            HEADING_EXTRACT_PROMPT.render(headings=headings), return_json=True
        )
        metadata = []
        sections = []
        tasks = []

        async with asyncio.TaskGroup() as tg:
            for chunk in split_markdown_by_headings(
                markdown_content, headings, adjusted_headings
            ):
                task1 = tg.create_task(
                    cls._parse_chunk_async(
                        doc_extractor,
                        language_model,
                        vision_model,
                        table_model,
                        None,
                        chunk,
                        image_dir,
                    )
                )
                task2 = tg.create_task(
                    language_model(
                        SECTION_SUMMARY_PROMPT.render(section_content=chunk),
                    )
                )
                tasks.append((task1, task2))

        # Process results in order
        for task1, task2 in tasks:
            meta, section = task1.result()
            metadata.append(meta)
            sections.append(section)
            for section in sections:
                section.summary = task2.result()

        merged_metadata = await language_model(
            MERGE_METADATA_PROMPT.render(metadata=metadata, topic=topic),
            return_json=True,
        )

        # 确保merged_metadata是字典类型
        if not isinstance(merged_metadata, dict):
            logger.warning(f"merged_metadata is not a dict: {type(merged_metadata)}, using default metadata")
            merged_metadata = {
                "title": "未知标题",
                "author": "未知作者",
                "subject": "演示文稿"
            }

        return Document(
            image_dir=image_dir, metadata=merged_metadata, sections=sections
        )

    def __contains__(self, key: str):
        for section in self.sections:
            if section.title == key:
                return True
        return False

    def __getitem__(self, key: str):
        for section in self.sections:
            if section.title == key:
                return section
        raise KeyError(
            f"section not found: {key}, available sections: {[section.title for section in self.sections]}"
        )

    def to_dict(self):
        return asdict(self)

    def retrieve(
        self,
        indexs: dict[str, list[str]],
    ) -> list[SubSection]:
        assert isinstance(
            indexs, dict
        ), "subsection_keys for index must be a dict, follow a two-level structure"
        subsecs = []
        for sec_key, subsec_keys in indexs.items():
            section = self[sec_key]
            for subsec_key in subsec_keys:
                subsecs.append(section[subsec_key])
        return subsecs

    def find_caption(self, caption: str):
        for media in self.iter_medias():
            if media.caption == caption:
                return media.path
        raise ValueError(f"Image caption not found: {caption}")

    def get_overview(self, include_summary: bool = False):
        overview = ""
        for section in self.sections:
            overview += f"Section: {section.title}\n"
            if include_summary:
                overview += f"\tSummary: {section.summary}\n"
            for subsection in section.subsections:
                overview += f"\tSubsection: {subsection.title}\n"
                for media in subsection.medias:
                    overview += f"\t\tMedia: {media.caption}\n"
                overview += "\n"
        return overview

    @property
    def metainfo(self):
        return "\n".join([f"{k}: {v}" for k, v in self.metadata.items()])

    @property
    def subsections(self):
        return [subsec for section in self.sections for subsec in section.subsections]


@dataclass
class OutlineItem:
    purpose: str
    section: str
    indexs: dict[str, list[str]] | str
    images: list[str]

    @classmethod
    def from_dict(cls, data: dict[str, Any]):
        assert (
            "purpose" in data and "section" in data
        ), "purpose and section of outline item are required"
        return cls(
            purpose=data["purpose"],
            section=data["section"],
            indexs=data.get("indexs", {}),
            images=data.get("images", []),
        )

    def retrieve(self, slide_idx: int, document: Document):
        subsections = document.retrieve(self.indexs)
        header = f"Slide-{slide_idx+1}: {self.purpose}\n"
        content = ""
        for subsection in subsections:
            content += f"Paragraph: {subsection.title}\nContent: {subsection.content}\n"
        images = [
            f"Image: {document.find_caption(caption)}\nCaption: {caption}"
            for caption in self.images
        ]
        return header, content, images

    def check_retrieve(self, document: Document, sim_bound: float):
        for sec_key, subsec_keys in list(self.indexs.items()):
            section = max(
                document.sections, key=lambda x: edit_distance(x.title, sec_key)
            )
            self.indexs[section.title] = self.indexs.pop(sec_key)
            if edit_distance(section.title, sec_key) < sim_bound:
                logger.warning(
                    f"section not found: {sec_key}, available sections: {[section.title for section in document.sections]}.",
                )
                raise ValueError(
                    f"section not found: {sec_key}, available sections: {[section.title for section in document.sections]}."
                )
            for idx in range(len(subsec_keys)):
                subsection = max(
                    section.subsections,
                    key=lambda x: edit_distance(x.title, subsec_keys[idx]),
                )
                self.indexs[section.title][idx] = subsection.title
                if edit_distance(subsection.title, subsec_keys[idx]) < sim_bound:
                    raise ValueError(
                        f"subsection {subsec_keys[idx]} not found in section {section.title}, available subsections: {[subsection.title for subsection in section.subsections]}."
                    )

    def check_images(self, document: Document, text_model: LLM, sim_bound: float):
        doc_images = list(document.iter_medias())
        image_embeddings = []
        for idx, image in enumerate(self.images):
            if len(doc_images) == 0:
                raise ValueError("Document does not contain any images.")
            similar = max(doc_images, key=lambda x: edit_distance(x.caption, image))
            if edit_distance(similar.caption, image) > sim_bound:
                self.images[idx] = similar.caption
                continue
            if len(image_embeddings) == 0:
                image_embeddings.extend(
                    [text_model.get_embedding(image) for image in self.images]
                )

            embedding = text_model.get_embedding(image)
            similar = max(
                range(len(image_embeddings)),
                key=lambda x: cosine_similarity(embedding, image_embeddings[x]),
            )
            if cosine_similarity(embedding, image_embeddings[similar]) > sim_bound:
                self.images[idx] = doc_images[similar].caption
            else:
                logger.warning(
                    f"image not found: {image}, available images: {[image.caption for image in doc_images]}.",
                )
                raise ValueError(
                    f"image not found: {image}, available images: \n{[image.caption for image in doc_images]}\nPlease ensure the caption is exactly matched."
                )

    async def check_images_async(
        self, document: Document, text_model: AsyncLLM, sim_bound: float
    ):
        doc_images = list(document.iter_medias())
        image_embeddings = []
        for idx, image in enumerate(self.images):
            if len(doc_images) == 0:
                raise ValueError("Document does not contain any images.")
            similar = max(doc_images, key=lambda x: edit_distance(x.caption, image))
            if edit_distance(similar.caption, image) > sim_bound:
                self.images[idx] = similar.caption
                continue
            if len(image_embeddings) == 0:
                image_embeddings = await asyncio.gather(
                    *[text_model.get_embedding(image) for image in self.images]
                )

            embedding = await text_model.get_embedding(image)
            similar = max(
                range(len(image_embeddings)),
                key=lambda x: cosine_similarity(embedding, image_embeddings[x]),
            )
            if cosine_similarity(embedding, image_embeddings[similar]) > sim_bound:
                self.images[idx] = doc_images[similar].caption
